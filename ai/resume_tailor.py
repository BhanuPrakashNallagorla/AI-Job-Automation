"""
Resume Tailor using Claude Opus 4.
Intelligently tailors resumes to job descriptions while maintaining truthfulness.
"""
import json
import os
import difflib
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
import structlog

from anthropic import Anthropic
from docx import Document
from docx.shared import Inches, Pt
from tenacity import retry, stop_after_attempt, wait_exponential

from config import settings


logger = structlog.get_logger(__name__)


class TailoringLevel:
    """Tailoring intensity levels."""
    CONSERVATIVE = "conservative"
    MODERATE = "moderate"
    AGGRESSIVE = "aggressive"


class ResumeTailor:
    """
    Tailors resumes to job descriptions using Claude Opus 4.
    
    Features:
    - Three tailoring levels (conservative, moderate, aggressive)
    - Maintains truthfulness - no fabrication
    - ATS-friendly output
    - Detailed change tracking
    - DOCX output generation
    - Match score calculation
    """
    
    # ========================================================================
    # Prompt Templates
    # ========================================================================
    
    TAILORING_PROMPTS = {
        TailoringLevel.CONSERVATIVE: """You are an expert resume consultant. Your task is to make MINIMAL, SUBTLE adjustments to tailor this resume for the target job.

CONSERVATIVE tailoring guidelines:
- Only reorder existing bullet points to put most relevant first
- Add 1-2 keywords from job description ONLY if they naturally fit
- Adjust the professional summary slightly
- DO NOT add new skills or experiences
- DO NOT modify job titles or dates
- DO NOT exaggerate accomplishments

Be very careful to maintain complete truthfulness.""",

        TailoringLevel.MODERATE: """You are an expert resume consultant. Your task is to make MODERATE adjustments to tailor this resume for the target job.

MODERATE tailoring guidelines:
- Reorder bullet points to emphasize relevant experience
- Add relevant keywords naturally throughout
- Expand relevant bullet points with more detail
- Adjust the professional summary to highlight matching qualifications
- Highlight transferable skills that match job requirements
- DO NOT fabricate experiences or skills
- DO NOT exaggerate beyond reasonable interpretation

The resume should feel optimized but remain truthful.""",

        TailoringLevel.AGGRESSIVE: """You are an expert resume consultant. Your task is to SIGNIFICANTLY restructure this resume for the target job while maintaining truthfulness.

AGGRESSIVE tailoring guidelines:
- Completely restructure to lead with most relevant experience
- Rewrite bullet points to directly mirror job description language
- Add a "Key Skills" or "Core Competencies" section if not present
- Significantly expand relevant sections, condense less relevant ones
- Use exact keywords from job description wherever truthfully applicable
- Consider combining or separating sections for better flow
- Make the summary highly targeted to this specific role
- STILL maintain complete truthfulness - never fabricate

The resume should feel custom-built for this job."""
    }
    
    MAIN_PROMPT = """{level_prompt}

<current_resume>
{resume_text}
</current_resume>

<job_description>
{job_description}
</job_description>

<job_analysis>
{job_analysis}
</job_analysis>

Please tailor the resume and provide your response in the following JSON format:

{{
    "tailored_resume": {{
        "contact_info": {{
            "name": "candidate name",
            "email": "email",
            "phone": "phone",
            "location": "city, state",
            "linkedin": "linkedin url if available"
        }},
        "professional_summary": "tailored professional summary paragraph",
        "skills": {{
            "technical": ["skill1", "skill2"],
            "frameworks": ["framework1"],
            "tools": ["tool1"],
            "soft_skills": ["skill1"]
        }},
        "experience": [
            {{
                "job_title": "title",
                "company": "company name",
                "location": "location",
                "start_date": "start",
                "end_date": "end or Present",
                "bullets": [
                    "accomplishment 1",
                    "accomplishment 2"
                ]
            }}
        ],
        "education": [
            {{
                "degree": "degree name",
                "institution": "school name",
                "year": "graduation year",
                "gpa": "gpa if mentioned",
                "relevant_coursework": ["course1"]
            }}
        ],
        "projects": [
            {{
                "name": "project name",
                "description": "brief description",
                "technologies": ["tech1"],
                "highlights": ["highlight1"]
            }}
        ],
        "certifications": ["cert1", "cert2"],
        "additional_sections": {{}}
    }},
    "changes_made": [
        {{
            "section": "section name",
            "change_type": "reorder|modify|add|remove|expand",
            "description": "what was changed and why",
            "original": "original text if applicable",
            "modified": "new text if applicable"
        }}
    ],
    "keywords_added": ["keyword1", "keyword2"],
    "keywords_matched": ["existing keywords that match job"],
    "ats_optimization": {{
        "score": <0-100>,
        "improvements": ["improvement made 1"],
        "remaining_gaps": ["gap that couldn't be addressed"]
    }},
    "truthfulness_check": {{
        "all_claims_accurate": true,
        "notes": "any notes about maintaining truthfulness"
    }}
}}

Important:
1. Maintain ALL original information - just restructure and enhance
2. Every change must be documented in changes_made
3. Keywords should be added naturally, not stuffed
4. Be honest in ats_optimization about what gaps remain
5. Respond ONLY with valid JSON"""

    def __init__(self, api_key: Optional[str] = None):
        """Initialize with Anthropic API key."""
        self.api_key = api_key or settings.anthropic_api_key
        self.client = Anthropic(api_key=self.api_key) if self.api_key else None
        self.model = settings.claude_opus_model
        self.output_dir = Path(settings.resumes_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.logger = logger.bind(component="ResumeTailor")
    
    def read_docx(self, file_path: str) -> str:
        """Read text content from a DOCX file."""
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    
    def create_docx(
        self,
        resume_data: Dict[str, Any],
        output_path: str,
    ) -> str:
        """Create a formatted DOCX resume from structured data."""
        doc = Document()
        
        # Set narrow margins for more content
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # Contact Info Header
        contact = resume_data.get("contact_info", {})
        if contact.get("name"):
            name_para = doc.add_heading(contact["name"], level=0)
            name_para.alignment = 1  # Center
        
        contact_line = []
        if contact.get("email"):
            contact_line.append(contact["email"])
        if contact.get("phone"):
            contact_line.append(contact["phone"])
        if contact.get("location"):
            contact_line.append(contact["location"])
        if contact.get("linkedin"):
            contact_line.append(contact["linkedin"])
        
        if contact_line:
            contact_para = doc.add_paragraph(" | ".join(contact_line))
            contact_para.alignment = 1
        
        # Professional Summary
        if resume_data.get("professional_summary"):
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(resume_data["professional_summary"])
        
        # Skills
        skills = resume_data.get("skills", {})
        if any(skills.values()):
            doc.add_heading("Skills", level=1)
            for category, skill_list in skills.items():
                if skill_list:
                    category_name = category.replace("_", " ").title()
                    doc.add_paragraph(
                        f"{category_name}: {', '.join(skill_list)}",
                        style="List Bullet"
                    )
        
        # Experience
        experience = resume_data.get("experience", [])
        if experience:
            doc.add_heading("Experience", level=1)
            for job in experience:
                # Job title and company
                job_para = doc.add_paragraph()
                job_para.add_run(job.get("job_title", "")).bold = True
                job_para.add_run(f" | {job.get('company', '')}")
                
                # Dates and location
                date_location = []
                if job.get("location"):
                    date_location.append(job["location"])
                date_range = f"{job.get('start_date', '')} - {job.get('end_date', 'Present')}"
                date_location.append(date_range)
                doc.add_paragraph(" | ".join(date_location))
                
                # Bullet points
                for bullet in job.get("bullets", []):
                    doc.add_paragraph(bullet, style="List Bullet")
        
        # Education
        education = resume_data.get("education", [])
        if education:
            doc.add_heading("Education", level=1)
            for edu in education:
                edu_para = doc.add_paragraph()
                edu_para.add_run(edu.get("degree", "")).bold = True
                edu_para.add_run(f" | {edu.get('institution', '')}")
                if edu.get("year"):
                    edu_para.add_run(f" | {edu['year']}")
                if edu.get("gpa"):
                    doc.add_paragraph(f"GPA: {edu['gpa']}")
        
        # Projects
        projects = resume_data.get("projects", [])
        if projects:
            doc.add_heading("Projects", level=1)
            for project in projects:
                project_para = doc.add_paragraph()
                project_para.add_run(project.get("name", "")).bold = True
                if project.get("technologies"):
                    project_para.add_run(f" ({', '.join(project['technologies'])})")
                if project.get("description"):
                    doc.add_paragraph(project["description"])
                for highlight in project.get("highlights", []):
                    doc.add_paragraph(highlight, style="List Bullet")
        
        # Certifications
        certs = resume_data.get("certifications", [])
        if certs:
            doc.add_heading("Certifications", level=1)
            for cert in certs:
                doc.add_paragraph(cert, style="List Bullet")
        
        # Save
        doc.save(output_path)
        return output_path
    
    def generate_diff(
        self,
        original_text: str,
        modified_text: str
    ) -> List[Dict[str, Any]]:
        """Generate a human-readable diff of changes."""
        original_lines = original_text.split("\n")
        modified_lines = modified_text.split("\n")
        
        differ = difflib.unified_diff(
            original_lines,
            modified_lines,
            lineterm="",
            n=0
        )
        
        changes = []
        for line in differ:
            if line.startswith("---") or line.startswith("+++") or line.startswith("@@"):
                continue
            if line.startswith("-"):
                changes.append({"type": "removed", "text": line[1:]})
            elif line.startswith("+"):
                changes.append({"type": "added", "text": line[1:]})
        
        return changes
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10)
    )
    async def tailor(
        self,
        resume_path: str,
        job_description: str,
        job_analysis: Optional[Dict[str, Any]] = None,
        tailoring_level: str = TailoringLevel.MODERATE,
        output_filename: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Tailor a resume to a job description.
        
        Args:
            resume_path: Path to the original resume (DOCX)
            job_description: Full job description text
            job_analysis: Optional pre-analyzed JD (from JDAnalyzer)
            tailoring_level: conservative, moderate, or aggressive
            output_filename: Custom output filename
            
        Returns:
            Dict with tailored resume data, file path, changes, and match score
        """
        if not self.client:
            raise ValueError("Anthropic API key not configured")
        
        self.logger.info(
            "Starting resume tailoring",
            level=tailoring_level,
            resume=resume_path
        )
        
        # Read original resume
        resume_text = self.read_docx(resume_path)
        
        # Get job analysis if not provided
        if not job_analysis:
            from ai.jd_analyzer import JDAnalyzer
            analyzer = JDAnalyzer(api_key=self.api_key)
            job_analysis = await analyzer.analyze(job_description)
        
        # Get the appropriate prompt
        level_prompt = self.TAILORING_PROMPTS.get(
            tailoring_level, 
            self.TAILORING_PROMPTS[TailoringLevel.MODERATE]
        )
        
        # Build the full prompt
        prompt = self.MAIN_PROMPT.format(
            level_prompt=level_prompt,
            resume_text=resume_text,
            job_description=job_description,
            job_analysis=json.dumps(job_analysis, indent=2)
        )
        
        try:
            # Call Claude Opus
            message = self.client.messages.create(
                model=self.model,
                max_tokens=settings.max_tokens_resume_tailor,
                messages=[{
                    "role": "user",
                    "content": prompt
                }]
            )
            
            # Track usage
            input_tokens = message.usage.input_tokens
            output_tokens = message.usage.output_tokens
            
            self.logger.info(
                "Resume tailoring complete",
                input_tokens=input_tokens,
                output_tokens=output_tokens
            )
            
            # Parse response
            response_text = message.content[0].text
            if response_text.startswith("```"):
                lines = response_text.split("\n")
                response_text = "\n".join(lines[1:-1])
            
            result = json.loads(response_text)
            
            # Generate output filename
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                base_name = Path(resume_path).stem
                output_filename = f"{base_name}_tailored_{tailoring_level}_{timestamp}.docx"
            
            output_path = self.output_dir / output_filename
            
            # Create the tailored DOCX
            self.create_docx(result["tailored_resume"], str(output_path))
            
            # Calculate match score
            ats_score = result.get("ats_optimization", {}).get("score", 0)
            
            return {
                "success": True,
                "output_path": str(output_path),
                "tailored_resume": result["tailored_resume"],
                "changes_made": result.get("changes_made", []),
                "keywords_added": result.get("keywords_added", []),
                "keywords_matched": result.get("keywords_matched", []),
                "ats_score": ats_score,
                "ats_improvements": result.get("ats_optimization", {}).get("improvements", []),
                "remaining_gaps": result.get("ats_optimization", {}).get("remaining_gaps", []),
                "truthfulness_verified": result.get("truthfulness_check", {}).get("all_claims_accurate", True),
                "tailoring_level": tailoring_level,
                "metadata": {
                    "model": self.model,
                    "input_tokens": input_tokens,
                    "output_tokens": output_tokens,
                }
            }
            
        except json.JSONDecodeError as e:
            self.logger.error("Failed to parse tailoring response", error=str(e))
            raise ValueError(f"Invalid response from AI: {e}")
        except Exception as e:
            self.logger.error("Resume tailoring failed", error=str(e))
            raise
    
    async def generate_all_versions(
        self,
        resume_path: str,
        job_description: str,
        job_analysis: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Dict[str, Any]]:
        """
        Generate all three tailoring versions.
        
        Returns:
            Dict with conservative, moderate, and aggressive versions
        """
        results = {}
        
        for level in [TailoringLevel.CONSERVATIVE, TailoringLevel.MODERATE, TailoringLevel.AGGRESSIVE]:
            self.logger.info(f"Generating {level} version")
            try:
                result = await self.tailor(
                    resume_path=resume_path,
                    job_description=job_description,
                    job_analysis=job_analysis,
                    tailoring_level=level,
                )
                results[level] = result
            except Exception as e:
                self.logger.error(f"Failed to generate {level} version", error=str(e))
                results[level] = {"success": False, "error": str(e)}
        
        return results


# ============================================================================
# Example Prompts That Work Well
# ============================================================================

EXAMPLE_PROMPTS = {
    "software_engineer": """Tailor this resume for a Senior Software Engineer role.
Focus on:
- Scalability and system design experience
- Production deployment and DevOps practices
- Leadership and mentoring
- Metrics and impact quantification""",

    "data_scientist": """Tailor this resume for a Data Scientist role.
Focus on:
- ML model development and deployment
- Statistical analysis and experimentation
- Python/R programming proficiency
- Business impact of data projects""",

    "product_manager": """Tailor this resume for a Product Manager role.
Focus on:
- Product strategy and roadmap planning
- Cross-functional collaboration
- User research and metrics
- Shipping products and driving growth"""
}


# ============================================================================
# Convenience Function
# ============================================================================

async def tailor_resume(
    resume_path: str,
    job_description: str,
    tailoring_level: str = TailoringLevel.MODERATE,
    api_key: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Convenience function to tailor a resume.
    
    Example:
        result = await tailor_resume(
            resume_path="resumes/my_resume.docx",
            job_description="We are looking for a Python developer...",
            tailoring_level="moderate"
        )
        print(f"Tailored resume saved to: {result['output_path']}")
        print(f"ATS Score: {result['ats_score']}")
    """
    tailor = ResumeTailor(api_key=api_key)
    return await tailor.tailor(
        resume_path=resume_path,
        job_description=job_description,
        tailoring_level=tailoring_level,
    )
