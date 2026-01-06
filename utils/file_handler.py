"""
File Handler Utilities.
Manages resume uploads, DOCX operations, and file sanitization.
"""
import os
import re
import shutil
from pathlib import Path
from typing import Optional, Tuple, List
from datetime import datetime
import structlog

from docx import Document
import aiofiles

from config import settings


logger = structlog.get_logger(__name__)


class FileHandler:
    """
    Handles file operations for resumes and cover letters.
    
    Features:
    - Resume upload validation
    - DOCX read/write
    - File sanitization
    - Directory management
    """
    
    ALLOWED_EXTENSIONS = {".docx", ".doc", ".pdf", ".txt"}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
    
    def __init__(self):
        """Initialize with configured directories."""
        self.resumes_dir = Path(settings.resumes_dir)
        self.cover_letters_dir = Path(settings.cover_letters_dir)
        self.scraped_jobs_dir = Path(settings.scraped_jobs_dir)
        
        # Ensure directories exist
        self.create_directories()
        
        self.logger = logger.bind(component="FileHandler")
    
    def create_directories(self) -> None:
        """Create required directories if they don't exist."""
        for dir_path in [self.resumes_dir, self.cover_letters_dir, self.scraped_jobs_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)
    
    def validate_file(
        self, 
        file_path: str, 
        allowed_extensions: Optional[set] = None
    ) -> Tuple[bool, str]:
        """
        Validate a file for upload.
        
        Args:
            file_path: Path to the file
            allowed_extensions: Set of allowed extensions (default: ALLOWED_EXTENSIONS)
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        path = Path(file_path)
        
        if not path.exists():
            return False, f"File not found: {file_path}"
        
        # Check extension
        extensions = allowed_extensions or self.ALLOWED_EXTENSIONS
        if path.suffix.lower() not in extensions:
            return False, f"Invalid file type. Allowed: {', '.join(extensions)}"
        
        # Check file size
        file_size = path.stat().st_size
        if file_size > self.MAX_FILE_SIZE:
            return False, f"File too large. Maximum size: {self.MAX_FILE_SIZE // (1024*1024)} MB"
        
        if file_size == 0:
            return False, "File is empty"
        
        return True, ""
    
    async def handle_resume_upload(
        self,
        file_content: bytes,
        original_filename: str,
        is_base_resume: bool = False,
    ) -> Tuple[str, dict]:
        """
        Handle resume file upload.
        
        Args:
            file_content: File content as bytes
            original_filename: Original filename
            is_base_resume: Whether this is the base resume template
            
        Returns:
            Tuple of (saved_path, metadata)
        """
        # Sanitize filename
        safe_filename = sanitize_filename(original_filename)
        
        # Add timestamp for uniqueness
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        name, ext = os.path.splitext(safe_filename)
        
        if is_base_resume:
            final_filename = f"base_resume_{timestamp}{ext}"
        else:
            final_filename = f"{name}_{timestamp}{ext}"
        
        save_path = self.resumes_dir / final_filename
        
        # Save file
        async with aiofiles.open(save_path, "wb") as f:
            await f.write(file_content)
        
        self.logger.info(
            "Resume uploaded",
            path=str(save_path),
            size=len(file_content),
            is_base=is_base_resume
        )
        
        metadata = {
            "original_filename": original_filename,
            "saved_filename": final_filename,
            "file_path": str(save_path),
            "file_size": len(file_content),
            "uploaded_at": datetime.now().isoformat(),
            "is_base_resume": is_base_resume,
        }
        
        return str(save_path), metadata
    
    def list_resumes(self) -> List[dict]:
        """List all resumes in the resumes directory."""
        resumes = []
        
        for file_path in self.resumes_dir.glob("*"):
            if file_path.suffix.lower() in self.ALLOWED_EXTENSIONS:
                resumes.append({
                    "filename": file_path.name,
                    "path": str(file_path),
                    "size": file_path.stat().st_size,
                    "modified": datetime.fromtimestamp(
                        file_path.stat().st_mtime
                    ).isoformat(),
                    "is_base": "base_resume" in file_path.name,
                })
        
        return sorted(resumes, key=lambda x: x["modified"], reverse=True)
    
    def delete_file(self, file_path: str) -> bool:
        """Delete a file safely."""
        path = Path(file_path)
        
        # Security check - only delete from our directories
        if not (
            str(path).startswith(str(self.resumes_dir)) or
            str(path).startswith(str(self.cover_letters_dir))
        ):
            self.logger.warning("Attempted to delete file outside allowed directories", path=file_path)
            return False
        
        if path.exists():
            path.unlink()
            self.logger.info("File deleted", path=file_path)
            return True
        
        return False


def read_docx(file_path: str) -> str:
    """
    Read text content from a DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    doc = Document(file_path)
    full_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)


def save_docx(content: str, file_path: str) -> str:
    """
    Save text content to a DOCX file.
    
    Args:
        content: Text content to save
        file_path: Output file path
        
    Returns:
        Saved file path
    """
    doc = Document()
    
    # Split content by paragraphs and add to document
    for paragraph in content.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    
    doc.save(file_path)
    return file_path


def sanitize_filename(filename: str) -> str:
    """
    Sanitize a filename to remove unsafe characters.
    
    Args:
        filename: Original filename
        
    Returns:
        Sanitized filename
    """
    # Remove path components
    filename = os.path.basename(filename)
    
    # Remove or replace unsafe characters
    # Keep only alphanumeric, dash, underscore, and dot
    safe_chars = re.sub(r"[^\w\-_. ]", "", filename)
    
    # Replace spaces with underscores
    safe_chars = safe_chars.replace(" ", "_")
    
    # Remove multiple underscores/dots
    safe_chars = re.sub(r"_+", "_", safe_chars)
    safe_chars = re.sub(r"\.+", ".", safe_chars)
    
    # Limit length
    name, ext = os.path.splitext(safe_chars)
    if len(name) > 100:
        name = name[:100]
    
    return f"{name}{ext}"


def get_file_info(file_path: str) -> dict:
    """Get file information."""
    path = Path(file_path)
    
    if not path.exists():
        return {"exists": False}
    
    stat = path.stat()
    
    return {
        "exists": True,
        "name": path.name,
        "path": str(path.absolute()),
        "size": stat.st_size,
        "size_human": format_file_size(stat.st_size),
        "extension": path.suffix.lower(),
        "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
    }


def format_file_size(size_bytes: int) -> str:
    """Format file size in human-readable format."""
    for unit in ["B", "KB", "MB", "GB"]:
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"
